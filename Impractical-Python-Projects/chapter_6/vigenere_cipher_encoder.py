#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Jan 14 12:16:07 2019

@author: toddbilsborough

Project #12: Hiding a Vigenere Cipher project
from Impractical Python Projects

Objective
- Hide a secret message summarizing bid details within an 
official-looking text document. Start with an unencrypted message
and finish with an encrypted version
- The message is going to an intermediary and then to a competing
corporation.
- The message needs to evade email filters
- The message needs to be invisible to visual inspection

Notes
- The cipher uses a keyword and a table with all the letters on both axes. 
For each letter along the vertical axis, the alphabet starts at that letter
and cycles through the full alphabet (so the D row cycles through the 
alphabet from D to Z and then from A to C). The keyword is cycled over
the plaintext, and the letter in the keyword is the row to reference
that letter in the table.
- I'm not familiar with python-docx so I'll be referencing the book
more extensively for this one
- This project is going to use python-docx to hide the code
in a docx file. The book notes that there might be problems when
viewing the resulting doc in a non-word doc viewer (I'll be using
Google docs))
- The program will hide the message in the spaces between paragraphs
to avoid issues with kerning
- The message needs to have short lines to avoid them going over line
breaks in the final message

Process Notes
- Mostly just copying from the book to get a sense of the functions
and process
- The book first inserts the plaintext as a test. That works well enough.
- The book just has the cipher pre-generated in another file, which 
is weak, so I'm going to design my own encoder
- The final message should definitely be sent as a PDF. If it's loaded
in a word processor, the spellchecker immediately freaks out and underlines
all of the ciphertext, revealing its presence

Updates
- Updated program to check for number of blank lines, as per the chapter
6 practice project

"""

import docx
from docx.shared import RGBColor, Pt
import string
from itertools import cycle
import sys

def load_fake_text(file):
    """Return list of paragraph text from message"""
    text = docx.Document(file)
    return [p.text for p in text.paragraphs]

def load_real_text(file):
    """Return list of paragraph text from message, stripping
    blank lines"""
    text = docx.Document(file)
    return [p.text for p in text.paragraphs if len(p.text) > 0]

def verify_lines(fake_text, real_text):
    """Verifies that there are enough blank lines in the fake text
    to hold the message"""
    blank_line_count = fake_text.count("")
    real_line_count = len(real_text)
    if blank_line_count < real_line_count:
        print("\nInsufficient blank lines in fake message")
        print("Terminating")
        sys.exit()

def set_spacing(paragraph):
    """Use docx to set line spacing between paragraphs
    Copied directly from book"""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    
def generate_vigenere_table():
    """Generates a vigenere encoding table as a dict of dicts"""
    table = {}
    alpha = cycle(string.ascii_lowercase)
    for letter in string.ascii_lowercase:
        alpha_list = [next(alpha) for l in range(0, 26)]
        k_v_pairs = zip(string.ascii_lowercase, alpha_list)
        letter_dict = {i:a for (i, a) in k_v_pairs}
        table[letter] = letter_dict
        next(alpha) # Offset by 1 for the next letter
    return table

def encode_vigenere(plaintext_list, key):
    """Encodes plaintext into a vigenere cipher input. Spaces and 
    non-alphabetical characters do not move the key forward.
    input
    - plaintext (a list of strings)
    - key (a string of a single word)
    output
    - ciphertext as a list of strings"""
    # Turn the list into a single string
    plaintext = ""
    for line in plaintext_list:
        plaintext += "{}\n".format(line)
    key_cycle = cycle(list(key.lower()))
    table = generate_vigenere_table()
    ciphertext = []
    cipher_line = ""
    for c in plaintext:
        char = c.lower()
        if char == "\n":
            ciphertext.append(cipher_line)
            cipher_line = ""
            continue
        if char == " " or char not in string.ascii_lowercase:
            cipher_line += char
            continue
        if c.isupper():
            cipher_line += table[next(key_cycle)][char].upper()
            continue
        if c.islower():
            cipher_line += table[next(key_cycle)][char].lower()
            continue
    return ciphertext

def main():
    """Load the fake and real messages as lists.
    Turn the real message into a cipher.
    Add letterhead to new doc.
    Interleave the fake message and the cipher."""
    fake_list = load_fake_text('fakeMessage.docx')
    real_list = load_real_text('realMessage.docx')
    verify_lines(fake_list, real_list)
    key = "splendiferous"
    cipher_list = encode_vigenere(real_list, key)
    
    # Add letterhead to template
    doc = docx.Document('template.docx')
    doc.add_heading('Morland Holmes', 0) # 0 for max heading
    # 1 for subtitle heading
    subtitle = doc.add_heading('Global Consulting & Negotiations', 1)
    subtitle.alignment = 1
    doc.add_heading('', 1)
    doc.add_paragraph('January 14, 2019')
    doc.add_paragraph('')
    
    # interleave real and fake message lines
    # Copied from book from this point forward
    # With modifications to use my custom cipher encoding
    length_real = len(real_list)
    count_real = 0 # index of current line in hidden message
    for line in fake_list:
        if count_real < length_real and line == "":
            paragraph = doc.add_paragraph(cipher_list[count_real])
            paragraph_index = len(doc.paragraphs) - 1
            # set real message color to white
            run = doc.paragraphs[paragraph_index].runs[0]
            font = run.font
            font.color.rgb = RGBColor(255, 255, 255)
            count_real += 1
        else:
            paragraph = doc.add_paragraph(line)
        set_spacing(paragraph)
    doc.save('ciphertext_message_letterhead.docx')
    print("Done")
    
if __name__ == '__main__':
    main()
